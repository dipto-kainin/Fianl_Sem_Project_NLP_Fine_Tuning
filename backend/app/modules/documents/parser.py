"""
Documents Module - Multi-Format Document Parser.

Extracts text and metadata from PDF, DOCX, TXT, Markdown, HTML, and EPUB files.
Uses a strategy pattern to dispatch to the appropriate parser based on file type.
"""

import logging
from dataclasses import dataclass, field
from pathlib import Path

import chardet

logger = logging.getLogger(__name__)


@dataclass
class ParsedSection:
    """Represents a section of a parsed document."""
    title: str = ""
    content: str = ""
    level: int = 0  # Heading level (1=h1, 2=h2, etc.)
    page_numbers: list[int] = field(default_factory=list)


@dataclass
class ParsedDocument:
    """Result of parsing a document."""
    raw_text: str = ""
    sections: list[ParsedSection] = field(default_factory=list)
    metadata: dict = field(default_factory=dict)
    page_count: int = 0
    language: str | None = None


def parse_document(file_path: str, file_type: str) -> ParsedDocument:
    """
    Parse a document file and extract structured text.

    Args:
        file_path: Path to the document file.
        file_type: File extension (pdf, docx, txt, md, html, epub).

    Returns:
        ParsedDocument with extracted text, sections, and metadata.
    """
    parsers = {
        "pdf": _parse_pdf,
        "docx": _parse_docx,
        "txt": _parse_txt,
        "md": _parse_markdown,
        "html": _parse_html,
        "epub": _parse_epub,
    }

    parser_func = parsers.get(file_type.lower())
    if not parser_func:
        raise ValueError(f"Unsupported file type: {file_type}")

    logger.info(f"Parsing {file_type.upper()} document: {file_path}")
    result = parser_func(file_path)
    logger.info(
        f"Parsed document: {len(result.raw_text)} chars, "
        f"{len(result.sections)} sections, {result.page_count} pages"
    )
    return result


def _parse_pdf(file_path: str) -> ParsedDocument:
    """Parse a PDF document using PyMuPDF (fitz)."""
    import fitz  # PyMuPDF

    doc = fitz.open(file_path)
    sections: list[ParsedSection] = []
    full_text_parts: list[str] = []
    metadata = {}

    # Extract metadata
    pdf_meta = doc.metadata
    if pdf_meta:
        metadata = {
            "title": pdf_meta.get("title", ""),
            "author": pdf_meta.get("author", ""),
            "subject": pdf_meta.get("subject", ""),
            "creator": pdf_meta.get("creator", ""),
            "creation_date": pdf_meta.get("creationDate", ""),
        }
        # Remove empty values
        metadata = {k: v for k, v in metadata.items() if v}

    # Extract text page by page
    current_section = ParsedSection(title="Document Start", level=0, page_numbers=[])

    for page_num in range(len(doc)):
        page = doc[page_num]
        page_text = page.get_text("text")

        if page_text.strip():
            full_text_parts.append(page_text)

            # Try to detect headings from text blocks
            blocks = page.get_text("dict")["blocks"]
            for block in blocks:
                if block.get("type") == 0:  # Text block
                    for line in block.get("lines", []):
                        for span in line.get("spans", []):
                            text = span.get("text", "").strip()
                            font_size = span.get("size", 12)
                            is_bold = "bold" in span.get("font", "").lower()

                            # Heuristic: large or bold text is likely a heading
                            if text and font_size > 14 and is_bold:
                                # Save previous section if it has content
                                if current_section.content.strip():
                                    sections.append(current_section)
                                level = 1 if font_size > 18 else 2
                                current_section = ParsedSection(
                                    title=text,
                                    level=level,
                                    page_numbers=[page_num + 1],
                                )
                            elif text:
                                current_section.content += text + " "
                                if (page_num + 1) not in current_section.page_numbers:
                                    current_section.page_numbers.append(page_num + 1)

    # Add the last section
    if current_section.content.strip() or current_section.title:
        sections.append(current_section)

    page_count = len(doc)
    doc.close()

    # If no sections were detected, create one from the full text
    raw_text = "\n\n".join(full_text_parts)
    if not sections:
        sections = [ParsedSection(
            title="Full Document",
            content=raw_text,
            level=0,
            page_numbers=list(range(1, page_count + 1)),
        )]

    return ParsedDocument(
        raw_text=raw_text,
        sections=sections,
        metadata=metadata,
        page_count=page_count,
    )


def _parse_docx(file_path: str) -> ParsedDocument:
    """Parse a DOCX document using python-docx."""
    from docx import Document as DocxDocument

    doc = DocxDocument(file_path)
    sections: list[ParsedSection] = []
    full_text_parts: list[str] = []
    metadata = {}

    # Extract core properties
    core = doc.core_properties
    if core:
        metadata = {
            "title": core.title or "",
            "author": core.author or "",
            "subject": core.subject or "",
            "created": str(core.created) if core.created else "",
            "modified": str(core.modified) if core.modified else "",
        }
        metadata = {k: v for k, v in metadata.items() if v}

    # Map heading styles to levels
    heading_map = {
        "Heading 1": 1, "Heading 2": 2, "Heading 3": 3,
        "Heading 4": 4, "Heading 5": 5, "Heading 6": 6,
        "Title": 1,
    }

    current_section = ParsedSection(title="Document Start", level=0)

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        full_text_parts.append(text)
        style_name = para.style.name if para.style else ""

        if style_name in heading_map:
            # New heading found — save previous section
            if current_section.content.strip():
                sections.append(current_section)
            current_section = ParsedSection(
                title=text,
                level=heading_map[style_name],
            )
        else:
            current_section.content += text + "\n"

    # Add last section
    if current_section.content.strip() or current_section.title:
        sections.append(current_section)

    raw_text = "\n\n".join(full_text_parts)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                raw_text += "\n" + row_text

    return ParsedDocument(
        raw_text=raw_text,
        sections=sections if sections else [ParsedSection(title="Full Document", content=raw_text)],
        metadata=metadata,
        page_count=len(sections) or 1,
    )


def _parse_txt(file_path: str) -> ParsedDocument:
    """Parse a plain text file."""
    raw_bytes = Path(file_path).read_bytes()

    # Detect encoding
    detection = chardet.detect(raw_bytes)
    encoding = detection.get("encoding", "utf-8") or "utf-8"

    raw_text = raw_bytes.decode(encoding, errors="replace")

    # Split into paragraphs (double newlines)
    paragraphs = [p.strip() for p in raw_text.split("\n\n") if p.strip()]

    sections = [
        ParsedSection(title=f"Paragraph {i+1}", content=p, level=0)
        for i, p in enumerate(paragraphs)
    ]

    return ParsedDocument(
        raw_text=raw_text,
        sections=sections if sections else [ParsedSection(content=raw_text)],
        page_count=1,
    )


def _parse_markdown(file_path: str) -> ParsedDocument:
    """Parse a Markdown file, preserving heading structure."""
    import re

    raw_bytes = Path(file_path).read_bytes()
    detection = chardet.detect(raw_bytes)
    encoding = detection.get("encoding", "utf-8") or "utf-8"
    raw_text = raw_bytes.decode(encoding, errors="replace")

    sections: list[ParsedSection] = []
    current_section = ParsedSection(title="Document Start", level=0)

    # Match markdown headings
    heading_re = re.compile(r"^(#{1,6})\s+(.+)$", re.MULTILINE)

    lines = raw_text.split("\n")
    for line in lines:
        match = heading_re.match(line)
        if match:
            # Save previous section
            if current_section.content.strip() or current_section.title:
                sections.append(current_section)
            level = len(match.group(1))
            title = match.group(2).strip()
            current_section = ParsedSection(title=title, level=level)
        else:
            current_section.content += line + "\n"

    # Add last section
    if current_section.content.strip() or current_section.title:
        sections.append(current_section)

    return ParsedDocument(
        raw_text=raw_text,
        sections=sections if sections else [ParsedSection(content=raw_text)],
        page_count=1,
    )


def _parse_html(file_path: str) -> ParsedDocument:
    """Parse an HTML file using BeautifulSoup."""
    from bs4 import BeautifulSoup

    raw_bytes = Path(file_path).read_bytes()
    detection = chardet.detect(raw_bytes)
    encoding = detection.get("encoding", "utf-8") or "utf-8"
    html_content = raw_bytes.decode(encoding, errors="replace")

    soup = BeautifulSoup(html_content, "html.parser")

    # Extract metadata from <head>
    metadata = {}
    title_tag = soup.find("title")
    if title_tag:
        metadata["title"] = title_tag.get_text().strip()
    for meta in soup.find_all("meta"):
        name = meta.get("name", "").lower()
        content = meta.get("content", "")
        if name in ("author", "description", "keywords"):
            metadata[name] = content

    # Remove script and style elements
    for element in soup(["script", "style", "nav", "footer", "header"]):
        element.decompose()

    # Extract sections based on heading tags
    sections: list[ParsedSection] = []
    heading_tags = {"h1": 1, "h2": 2, "h3": 3, "h4": 4, "h5": 5, "h6": 6}

    current_section = ParsedSection(title="Document Start", level=0)

    for element in soup.find_all(True):
        if element.name in heading_tags:
            if current_section.content.strip():
                sections.append(current_section)
            current_section = ParsedSection(
                title=element.get_text().strip(),
                level=heading_tags[element.name],
            )
        elif element.name in ("p", "li", "td", "blockquote", "pre"):
            text = element.get_text().strip()
            if text:
                current_section.content += text + "\n"

    if current_section.content.strip() or current_section.title:
        sections.append(current_section)

    raw_text = soup.get_text(separator="\n", strip=True)

    return ParsedDocument(
        raw_text=raw_text,
        sections=sections if sections else [ParsedSection(content=raw_text)],
        metadata=metadata,
        page_count=1,
    )


def _parse_epub(file_path: str) -> ParsedDocument:
    """Parse an EPUB file using ebooklib + BeautifulSoup."""
    import ebooklib
    from ebooklib import epub
    from bs4 import BeautifulSoup

    book = epub.read_epub(file_path, options={"ignore_ncx": True})

    # Extract metadata
    metadata = {}
    title = book.get_metadata("DC", "title")
    if title:
        metadata["title"] = title[0][0]
    creator = book.get_metadata("DC", "creator")
    if creator:
        metadata["author"] = creator[0][0]
    language = book.get_metadata("DC", "language")
    lang_str = None
    if language:
        lang_str = language[0][0]
        metadata["language"] = lang_str

    # Extract text from each chapter
    sections: list[ParsedSection] = []
    full_text_parts: list[str] = []

    for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
        soup = BeautifulSoup(item.get_body_content(), "html.parser")

        # Try to get chapter title from first heading
        heading = soup.find(["h1", "h2", "h3"])
        chapter_title = heading.get_text().strip() if heading else item.get_name()

        chapter_text = soup.get_text(separator="\n", strip=True)
        if chapter_text.strip():
            full_text_parts.append(chapter_text)
            sections.append(ParsedSection(
                title=chapter_title,
                content=chapter_text,
                level=1,
            ))

    raw_text = "\n\n".join(full_text_parts)

    return ParsedDocument(
        raw_text=raw_text,
        sections=sections if sections else [ParsedSection(content=raw_text)],
        metadata=metadata,
        page_count=len(sections),
        language=lang_str,
    )
